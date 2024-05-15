from docx import Document
from docx.shared import Inches
from GameClass import Game
import ChessReader
import FigurePlotting
import StockfishGames
import OngoingGames




def createDocument(game):
    document = Document()
    document.add_heading('Chess game', 0)

    paragraph = document.add_paragraph('This document holds information about game number' + str(Game.GetChessEco(game)))
    #paragraph.add_run('bold').bold = True
    #paragraph.add_run(' and some ')
    #paragraph.add_run('italic.').italic = True

    document.add_heading('Section', level=1)
    document.add_paragraph('The following section contains information about the chosen chess game. This game was played by ' + Game.GetChessBlack(game) + 
                           ' as black player and ' + Game.GetChessWhite(game) + 'as white player. The score ended at ' + Game.GetChessResult(game), style='Intense Quote')

    document.add_heading('Subsection', level=2)
    document.add_paragraph('This the game was played ' + Game.GetChessDate(game) + '. The opening was a ' + Game.GetChessOpening(game)
                           + 'and the plycount ended at ' + Game.GetChessPlyCount(game), style='Intense Quote')
    
    document.add_paragraph('first item in unordered list', style='List Bullet')
    document.add_paragraph('first item in ordered list', style='List Number')



    totalWins = StockfishGames.totalWins()
    totalDrawn = StockfishGames.totalDrawn()
    totalLosses = StockfishGames.totalLost()
    whiteWins = StockfishGames.winsWhite()
    whiteDrawn = StockfishGames.drawnWhite()
    whiteLosses = StockfishGames.lostWhite()
    blackWins = StockfishGames.winsBlack()
    blackDrawn = StockfishGames.drawnBlack()
    blackLosses = StockfishGames.lostBlack()

    # Table for games won, drawn and lost for Stockfish
    records1 = (('Total', totalWins, totalDrawn, totalLosses),('White', whiteWins, whiteDrawn, whiteLosses),('Black', blackWins, blackDrawn, blackLosses))
    table1 = document.add_table(rows=1, cols=4, style="Table Grid")
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Colour'
    hdr_cells[1].text = 'Wins'
    hdr_cells[2].text = 'Drawn'
    hdr_cells[3].text = 'Lost'
    for colour, wins, drawn, lost in records1:
        row_cells = table1.add_row().cells
        row_cells[0].text = str(colour)
        row_cells[1].text = str(wins)
        row_cells[2].text = str(drawn)
        row_cells[3].text = str(lost)


    MD = OngoingGames.MeanDeviation('N')
    SD = OngoingGames.StandardDeviation('N')
    MDW = OngoingGames.MeanDeviation('Wh')
    SDW = OngoingGames.StandardDeviation('Wh')
    MDB = OngoingGames.MeanDeviation('B')
    SDB = OngoingGames.StandardDeviation('B')

    # Table for mean and standard deviations
    records2 = (('Mean, all games', MD),('Standard, all games', SD),('Mean, white games', MDW),('Standard, white games', SDW),('Mean, black games', MDB),('Standard, black games', SDB))
    table2 = document.add_table(rows=1, cols=2, style="Table Grid")
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = 'Type of deviation'
    hdr_cells[1].text = 'Value'
    for deviation, value in records2:
        row_cells = table2.add_row().cells
        row_cells[0].text = str(deviation)
        row_cells[1].text = str(value)

    document.add_page_break()
    document.save('Chess.docx')

    


gamelist = ChessReader.GetChessGameList("Stockfish.pgn")
game = gamelist[0]
createDocument(game)
print(game)
